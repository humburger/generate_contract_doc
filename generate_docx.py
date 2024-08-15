from docx import Document
from docx.shared import Cm
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from datetime import date

# 12pt; bold/normal/capslock; 
# heading 1 - 12pt bold; justify;
# normal text 12pt; justify; atkāpes (ja ir) un 'tab'
# page margins - 2 3 2 2 // approx

str_h1 = 'Heading 1'
str_body_text = 'Normal'
str_table_normal = 'Table Normal'
str_font_tnr = 'Times New Roman'
str_tab = '\t\t\t\t\t\t\t\t\t'

# page margins
CONST_TOP = 1.27
CONST_BOT = 1.27
CONST_LEFT = 3
CONST_RIGHT = 1

# paragraph indentation in inches
p_indentation = 0.49 

# document detination city and date
str_date_today = date.today()
str_date_year = str_date_today.strftime("%Y")
str_doc_dest_city = 'Rīgā'
str_doc_dest_date = fr'{str_date_year}. gada'

sec_party_org_name = ''
sec_party_org_addr = ''
sec_party_org_reg_num = ''
sec_party_contact_phone = ''
sec_party_contact_email  = ''
sec_party_contact_name = ''
sec_party_signature_name = ''
sec_party_signature_position = ''
sec_party_signature_dot_name = ''

# removes text before useable names which were received from input file and then prepares as normal string type value
def prepare_input(str_name):
    temp_arr = []
    
    temp_arr = str_name.split()
    temp_arr[0] = ""
    str_name = ' '.join(temp_arr).lstrip()
    
    return str_name

try:
    ## open read file for client data
    # value_name value
    # value_name2 value2 etc.
    f = open("contract details.txt", "r", encoding='utf-8')
    
    sec_party_org_name = prepare_input(f.readline().rstrip())
    sec_party_org_reg_num = prepare_input(f.readline().rstrip())
    sec_party_org_addr = prepare_input(f.readline().rstrip())
    
    f.readline() # reads empty line
    
    sec_party_contact_name = prepare_input(f.readline().rstrip())
    sec_party_contact_phone = prepare_input(f.readline().rstrip())
    sec_party_contact_email = prepare_input(f.readline().rstrip())
    
    f.readline() # reads empty line
    
    sec_party_signature_name = prepare_input(f.readline().rstrip())
    sec_party_signature_position = prepare_input(f.readline().rstrip())
    
    if sec_party_signature_name != '':
        # works on one surname
        sec_party_signature_dot_name = sec_party_signature_name[0] + '.' + sec_party_signature_name.split()[1]
        # sec_party_signature_dot_name = sec_party_signature_name.split()[1][0] + '.' + sec_party_signature_name.split()[2]
        
    ## close read file for client data
    f.close() 

except FileNotFoundError :
    print('Cannot find specified input file.')
    exit()

# first organisation's values
first_party_org_name = 'Marshmellow Alley'
first_party_addr = 'Pineapple Street 38'
first_party_phone = '123456789'
first_party_org_reg_num = '1234567891011'
first_party_position = 'CEO'
first_party_full_name = 'A.Wesker'

def page_margins(top, bot, lft, rght):
    # changing the page margins
    # https://stackoverflow.com/a/32916429
    sections = document.sections
    for s in sections:
        s.top_margin = Cm(top)
        s.bottom_margin = Cm(bot)
        s.left_margin = Cm(lft)
        s.right_margin = Cm(rght)

def change_def_style(str_name, sp_bef, sp_after, ln_spacing, font_name, font_size, font_color):
    # heading 1 style
    heading1_style = document.styles[str_name]
    heading1_style.paragraph_format.space_before = Pt(int(sp_bef))
    heading1_style.paragraph_format.space_after = Pt(int(sp_after))
    heading1_style.paragraph_format.line_spacing = float(ln_spacing)
    heading1_style.font.name = font_name
    heading1_style.font.size = Pt(int(font_size))
    heading1_style.font.color.rgb = font_color

def both_sides_properties_and_signatures():
    table = document.add_table(rows=6, cols=2)
    # hdr_cells = table.rows[0].cells
    table.rows[0].cells[0].text = f'{first_party_org_name}'
    table.rows[0].cells[1].text = f'{sec_party_org_name}'
    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
    
    # first org address
    # new line symbol after address and work position is for visual format only for data in invisible table
    table.rows[1].cells[0].text = f'Address: {first_party_addr}\n'
    table.rows[2].cells[0].text = f'Phone: {first_party_phone}'
    table.rows[3].cells[0].text = f'Registration number {first_party_org_reg_num}\n'
    table.rows[4].cells[0].text = f'{first_party_position}'
    table.rows[5].cells[0].text = f'{first_party_full_name}'

    # hdr_cells = table.rows[1].cells 
    # second org address
    table.rows[1].cells[1].text = f'Address: {sec_party_org_addr}'
    table.rows[2].cells[1].text = f'Phone: {sec_party_contact_phone}'
    table.rows[3].cells[1].text = f'Registration number {sec_party_org_reg_num}'
    table.rows[4].cells[1].text = f'{sec_party_signature_position}'
    table.rows[5].cells[1].text = f'{sec_party_signature_dot_name}'

    table.rows[1].cells[1].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    table.rows[2].cells[1].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    table.rows[3].cells[1].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    table.rows[4].cells[1].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    table.rows[5].cells[1].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

    # removes each rows unused spacing
    for r in table.rows: # takes each row
        for c  in r.cells: # takes each cell in row
            for p in c.paragraphs: # takes each paragraph's cell
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
        r.height = Cm(0.01)

def first_header1():
    heading = document.add_heading('Lorem Ipsum ', 1)
    yellow_marker(heading, '...')
    heading.add_run('\nLorem ipsum dolor sit amet,\n consectetur adipiscing elit.')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    # following line fixes that Heading 1 changes to "Times New Roman", because somehow stays "Calibri"
    # https://stackoverflow.com/a/60922725
    heading.style.element.rPr.rFonts.set(qn("w:asciiTheme"), "Times New Roman")

def new_header1(str_text):
    heading = document.add_heading(str_text.upper(), 1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    # following line fixes that Heading 1 changes to "Times New Roman", because somehow stays "Calibri"
    # https://stackoverflow.com/a/60922725
    heading.style.element.rPr.rFonts.set(qn("w:asciiTheme"), "Times New Roman")

def new_paragraph(str_text, indentation):
    p1 = document.add_paragraph(str(str_text))
    p1.paragraph_format.left_indent = Inches(indentation)
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def marked_paragraph1():
    p1 = document.add_paragraph(f'\n')
    p1.add_run('Marshmellow Alley').bold = True
    p1.add_run(', Nam id tellus porttitor, pharetra orci vitae, hendrerit diam.\n\n')

    # separate bold and marking with yellow marker
    sec_org_name_run = p1.add_run(f'{sec_party_org_name}')
    sec_org_name_run.bold = True
    sec_org_name_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # 

    p1.add_run(f', reģ. Nr. ')
    yellow_marker(p1, sec_party_org_reg_num)

    p1.add_run(f', kuru pārstāv ')
    yellow_marker(p1, sec_party_signature_name)

    p1.add_run(f', no otras puses,\n\n(abi kopā turpmāk – Puses, atsevišķi – Puse) nNam id tellus porttitor, pharetra orci vitae, hendrerit diam. (turpmāk – līgums):')
    # p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def marked_paragraph2():
    p = document.add_paragraph(f'1.2. Lorem ipsum dolor sit amet – ')
    yellow_marker(p, sec_party_org_name)
    p.add_run(f' darbinieki')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def marked_paragraph3(indentation):
    p = document.add_paragraph('2.5.2. Lorem ipsum dolor sit amet – ')
    yellow_marker(p, sec_party_contact_name)
    p.add_run(f', tālrunis ')
    yellow_marker(p, sec_party_contact_phone)
    p.add_run(f', elektroniskā pasta adrese: ')
    yellow_marker(p, sec_party_contact_email)
    p.paragraph_format.left_indent = Inches(indentation)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def contract_dest_date():
    destination = document.add_paragraph(f'{str_doc_dest_city}{str_tab}{str_doc_dest_date}')
    yellow_marker(destination, '...')

def end_paragraph():
    p1 = document.add_paragraph()
    end_marker_run = p1.add_run('ŠIS DOKUMENTS PARAKSTĪTS AR DROŠU ELEKTRONISKO PARAKSTU UN SATUR LAIKA ZĪMOGU')
    end_marker_run.font.size = Pt(9)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

# marks text with yellow marker
def yellow_marker(paragraph, str_text):
    paragraph.add_run(str_text).font.highlight_color = WD_COLOR_INDEX.YELLOW


### check if not imported into another script ###
if __name__ != "__main__": 
    raise ImportError('This module is not meant for importing')

"""
    main script part
"""
document = Document()

change_def_style(str_h1, 10, 10, 1, str_font_tnr, 12, None)
change_def_style(str_body_text, 0, 2, 1, str_font_tnr, 12, None)
page_margins(CONST_TOP, CONST_BOT, CONST_LEFT, CONST_RIGHT)

first_header1()
contract_dest_date()
marked_paragraph1()

new_header1('1. Lorem ipsum dolor sit amet')

new_paragraph('1.1. Pellentesque mollis vulputate pulvinar. Donec viverra commodo sapien, ac porta lorem mollis in. Phasellus tempus commodo elementum. ', 0)
marked_paragraph2()

new_header1('2. Lorem ipsum dolor sit amet')
new_paragraph('2.1. LPellentesque mollis vulputate pulvinar. Donec viverra commodo sapien, ac porta lorem mollis in. Phasellus tempus commodo elementum. ', 0)
new_paragraph('2.2. Pellentesque mollis vulputate pulvinar. Donec viverra commodo sapien, ac porta lorem mollis in. Phasellus tempus commodo elementum. ', 0)
new_paragraph('2.2.1. kLorem ipsum dolor sit amet, consectetur adipiscing elit. In consectetur lobortis finibus. Vestibulum volutpat mi quis massa pretium, porttitor maximus est venenatis', p_indentation)
new_paragraph('2.2.2. Pellentesque mollis vulputate pulvinar. Donec viverra commodo sapien, ac porta lorem mollis in. Phasellus tempus commodo elementum. .', p_indentation)
new_paragraph('2.3. Pellentesque mollis vulputate pulvinar. Donec viverra commodo sapien, ac porta lorem mollis in. Phasellus tempus commodo elementum. .', 0)
new_paragraph('2.4. Pellentesque mollis vulputate pulvinar. Donec viverra commodo sapien, ac porta lorem mollis in. Phasellus tempus commodo elementum. .', 0)
new_paragraph('2.5. Pellentesque mollis vulputate pulvinar. Donec viverra commodo sapien, ac porta lorem mollis in. Phasellus tempus commodo elementum. ', 0)
new_paragraph('2.5.1. Pellentesque mollis vulputate pulvinar. Donec viverra commodo sapien, ac porta lorem mollis in. Phasellus tempus commodo elementum. ', p_indentation)
marked_paragraph3(p_indentation)

new_header1('3. Lorem ipsum dolor sit amet')
new_paragraph('3.1. Lorem ipsum dolor sit amet, consectetur adipiscing elit. In consectetur lobortis finibus. Vestibulum volutpat mi quis massa pretium, porttitor maximus est venenatis', 0)
new_paragraph('3.2. Lorem ipsum dolor sit amet, consectetur adipiscing elit. In consectetur lobortis finibus. Vestibulum volutpat mi quis massa pretium, porttitor maximus est venenatis', 0)

new_header1('4. Lorem ipsum dolor sit amet')
new_paragraph('4.1. Lorem:', 0)
new_paragraph('4.1.1. Nullam finibus, nunc sit amet auctor tristique, elit dui vestibulum nunc, sed rutrum erat est quis odio. Suspendisse pharetra finibus lectus a dapibus.', p_indentation)
new_paragraph('4.1.2. Nullam finibus, nunc sit amet auctor tristique, elit dui vestibulum nunc, sed rutrum erat est quis odio. Suspendisse pharetra finibus lectus a dapibus.', p_indentation)
new_paragraph('4.1.3. Nullam finibus, nunc sit amet auctor tristique, elit dui vestibulum nunc, sed rutrum erat est quis odio. Suspendisse pharetra finibus lectus a dapibus.', p_indentation)
new_paragraph('4.1.4. Nullam finibus, nunc sit amet auctor tristique, elit dui vestibulum nunc, sed rutrum erat est quis odio. Suspendisse pharetra finibus lectus a dapibus.', p_indentation)
new_paragraph('4.2. Ipsum:', 0)
new_paragraph('4.2.1. Nullam finibus, nunc sit amet auctor tristique, elit dui vestibulum nunc, sed rutrum erat est quis odio. Suspendisse pharetra finibus lectus a dapibus.', p_indentation)
new_paragraph('4.2.2. Nullam finibus, nunc sit amet auctor tristique, elit dui vestibulum nunc, sed rutrum erat est quis odio. Suspendisse pharetra finibus lectus a dapibus.', p_indentation)
new_paragraph('4.2.3. Nullam finibus, nunc sit amet auctor tristique, elit dui vestibulum nunc, sed rutrum erat est quis odio. Suspendisse pharetra finibus lectus a dapibus.', p_indentation)

new_header1('5. Lorem ipsum dolor sit amet')

new_paragraph('5.1. Donec at aliquam nisl, vitae laoreet odio. Pellentesque consequat laoreet dapibus. Fusce rhoncus faucibus dui ut euismod. Vivamus et vestibulum augue.', 0)
new_paragraph('5.2. Donec at aliquam nisl, vitae laoreet odio. Pellentesque consequat laoreet dapibus. Fusce rhoncus faucibus dui ut euismod. Vivamus et vestibulum augue.', 0)
new_paragraph('5.3. Donec at aliquam nisl, vitae laoreet odio. Pellentesque consequat laoreet dapibus. Fusce rhoncus faucibus dui ut euismod. Vivamus et vestibulum augue.', 0)


new_header1('6. Lorem ipsum dolor sit amet')

new_paragraph('6.1. Vestibulum non mattis ex. Integer sed facilisis ligula. Mauris ullamcorper lacinia neque, non semper ex tincidunt quis. ', 0)
new_paragraph('6.2. Vestibulum non mattis ex. Integer sed facilisis ligula. Mauris ullamcorper lacinia neque, non semper ex tincidunt quis. ', 0)
new_paragraph('6.3. Vestibulum non mattis ex. Integer sed facilisis ligula. Mauris ullamcorper lacinia neque, non semper ex tincidunt quis. ', 0)
new_paragraph('6.4. Vestibulum non mattis ex. Integer sed facilisis ligula. Mauris ullamcorper lacinia neque, non semper ex tincidunt quis. ', 0)
new_paragraph('6.5. Vestibulum non mattis ex. Integer sed facilisis ligula. Mauris ullamcorper lacinia neque, non semper ex tincidunt quis. ', 0)

new_header1('7. Lorem ipsum dolor sit amet')
both_sides_properties_and_signatures()
new_paragraph('', 0)
end_paragraph()

document.save('compliance agreement.doc')

print("\nPlease revise generated document if there everything looks fine VISUALY and SEMANTICALLY (meaningful), i.e., \n\t > fonts can look weird (can solve by changing maually global text font to 'Times New Romans');\n\t > information in text makes sense where words were put by script;")